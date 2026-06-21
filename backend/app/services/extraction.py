import io
import logging

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    """Extract plain text from PDF, DOCX, or plain text/markdown files."""
    mime = mime_type.lower()

    if mime == "application/pdf":
        return _extract_pdf(file_bytes)
    elif mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_docx(file_bytes)
    elif mime in ("text/plain", "text/markdown", "text/x-markdown"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        # Attempt UTF-8 decode as fallback
        logger.warning("Unknown mime type %s — attempting UTF-8 decode", mime_type)
        return file_bytes.decode("utf-8", errors="replace")


def _extract_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n\n".join(paragraphs)
